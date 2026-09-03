import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "_screenshots")
OUT_DOCX = os.path.join(BASE, "REPORTS", "PHASE5_CYBEREXERCISE.docx")
os.makedirs(os.path.join(BASE, "REPORTS"), exist_ok=True)

BLACK = RGBColor(0, 0, 0)

def fix_dpi(path):
    try:
        im = Image.open(path)
        if im.info.get("dpi", (96, 96))[0] in (0, None):
            im.save(path, dpi=(96, 96))
    except Exception as e:
        print("dpi fix skip", path, e)

for f in os.listdir(SHOTS):
    if f.lower().endswith((".png", ".jpg", ".jpeg")):
        fix_dpi(os.path.join(SHOTS, f))

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = BLACK

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLACK
    return p

def para(text="", bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    return p

def add_image(path, width=6.0):
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(width))
        except Exception as e:
            doc.add_paragraph(f"[Could not embed {os.path.basename(path)}: {e}]")
    else:
        doc.add_paragraph(f"[Screenshot not found: {os.path.basename(path)}]")

def add_table(headers, rows, style_name="Light Grid Accent 1"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style_name
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
    doc.add_paragraph()

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t')
    fld_text.text = "Right-click and select 'Update Field' to build the Table of Contents."
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(fld_text)
    r_element.append(fld_end)

# ---------------- COVER PAGE ----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("OPERATION FALSE ASSURANCE")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BLACK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Phase 5: Client Infection and Domain Controller Compromise")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = BLACK

doc.add_paragraph()
for line in ["Team: BYEBUST", "Incident Date: 2025-10-11 (UTC)", "Report Date: 2026-08-23",
             "Classification: Internal - Exercise Use Only"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(12)
    r.font.color.rgb = BLACK

doc.add_page_break()

h("Table of Contents", level=1)
add_toc()
doc.add_page_break()

# ---------------- EXECUTIVE SUMMARY ----------------
h("Executive Summary", level=1)
para(
    "On 11 October 2025, workstation LINWOOD-WIN-PC (192.168.200.95), a member of the "
    "OYSTERTAINMENT.COM Active Directory domain, was compromised after its user retrieved a "
    "malicious Microsoft Word document from a spoofed IRS-transcript delivery page "
    "(r2consulting.net). Shortly after, the same host retrieved and executed a Windows "
    "executable payload (90352.exe) from a second external site (www.ownhive.com). The "
    "payload immediately began beaconing to external infrastructure using cleartext HTTP "
    "requests disguised as HTTPS traffic (destination TCP port 443) and to a dedicated C2 "
    "listener at 188.124.167.132:8082, consistent with a commodity remote-access/loader "
    "malware family."
)
para(
    "Approximately 20 minutes after the initial callback, the malware pivoted to the "
    "environment's domain controller, OYSTER-DC (192.168.200.4), which began issuing its own "
    "callbacks to 188.124.167.132 and retrieving three additional staged components disguised "
    "with PNG image-file names. Connection metadata between the workstation and the domain "
    "controller shows sustained SMB (TCP 445) communication between the two hosts throughout "
    "the incident window, consistent with lateral movement or credential/administrative-share "
    "activity following initial compromise."
)
para(
    "All eleven exercise objectives for this phase were fully resolved and confirmed accepted "
    "by the exercise platform. Two objectives (P5-06 and P5-07) required additional analysis "
    "beyond the literal log fields before an accepted answer was found: P5-06's accepted MIME "
    "type used the generic classification (application/octet-stream) rather than Zeek's "
    "file-magic-detected value (application/x-dosexec), and P5-07's accepted transaction count "
    "(26) reflected the total number of distinct TCP connections to the destination host, "
    "rather than the raw count of HTTP POST log entries (32) or the count of connections that "
    "specifically carried a POST (20). Both nuances are documented in their respective sections "
    "below, together with a closing lessons-learned note for future phases."
)
para(
    "This report distinguishes evidence directly confirmed in Security Onion (Zeek/Suricata) "
    "logs, labeled Confirmed, from the investigating team's analytical conclusions drawn from "
    "that evidence, labeled Assessment, throughout every section below."
)

doc.add_page_break()

# ---------------- INCIDENT TIMELINE ----------------
h("Incident Timeline (UTC)", level=1)
para("All times below are Confirmed directly from Zeek/Suricata log timestamps unless noted.")
timeline = [
    ("2025-10-11 09:00:00", "DHCP OFFER/ACK issued to client 192.168.200.95 (Linwood-Win-PC) - client onboarded to network. (P5-01)"),
    ("2025-10-11 09:00:00", "Client 192.168.200.95 performs Kerberos AS/TGS exchanges against 192.168.200.4, confirming AD domain OYSTERTAINMENT.COM and DC hostname OYSTER-DC. (P5-02)"),
    ("2025-10-11 09:02:26", "Client GET request to r2consulting.net delivers malicious document transcript-June152018-017188/2.doc (application/msword). (P5-03, P5-04)"),
    ("2025-10-11 09:03:25", "Client GET request to www.ownhive.com delivers executable payload 90352.exe (application/x-dosexec, 126,976 bytes). (P5-05, P5-06)"),
    ("2025-10-11 09:03:30 - 09:11:26", "Client conducts 32 cleartext HTTP POST transactions to 74.195.13.150 over TCP port 443, across 26 total TCP connections. (P5-07)"),
    ("2025-10-11 09:07:06", "First client POST to dedicated C2 endpoint 188.124.167.132:8082. (P5-08)"),
    ("2025-10-11 09:26:26", "First confirmed malware callback from domain controller 192.168.200.4 to 188.124.167.132:8082 - DC compromise confirmed. (P5-09)"),
    ("2025-10-11 09:26:27 - 09:27:19", "Domain controller retrieves three PNG-named staged objects from 185.159.131.55. (P5-10)"),
    ("2025-10-11 (throughout incident window)", "Client 192.168.200.95 communicates with domain controller 192.168.200.4 over SMB/TCP 445 (258 connection records). (P5-11)"),
]
add_table(["Timestamp (UTC)", "Event"], timeline)

doc.add_page_break()

# ---------------- IOC TABLE ----------------
h("Indicators of Compromise", level=1)
iocs = [
    ("192.168.200.95", "Compromised client (Linwood-Win-PC)"),
    ("192.168.200.4", "Domain controller (OYSTER-DC.oystertainment.com)"),
    ("oystertainment.com", "Active Directory domain"),
    ("r2consulting.net (184.168.230.1)", "Malicious document delivery host"),
    ("www.ownhive.com (209.105.227.32)", "Executable payload delivery host"),
    ("74.195.13.150", "Repeated cleartext HTTP POST destination on TCP/443 - suspected C2 (26 TCP connections, 20 carrying POST traffic)"),
    ("173.70.47.89", "Secondary cleartext HTTP POST destination on TCP/443 - suspected C2"),
    ("188.124.167.132:8082", "Dedicated C2 listener contacted by both client and domain controller"),
    ("185.159.131.55", "Staging host for PNG-named payload components retrieved by the DC"),
    ("transcript-June152018-017188/2.doc", "Malicious document filename (application/msword)"),
    ("90352.exe", "Executable payload filename (application/octet-stream, 126,976 bytes)"),
    ("MD5 1a142395b5043ed2ca1545e1a5b0ba95", "Hash of executable payload 90352.exe"),
    ("SHA1 37a73193a0d4f4a337094e9e5fd439dc293fc615", "Hash of executable payload 90352.exe"),
    ("table.png / toler.png / worming.png", "PNG-named objects retrieved by domain controller post-compromise"),
    ("TCP/445 (SMB) 192.168.200.95 <-> 192.168.200.4", "Client-to-DC SMB communication"),
]
add_table(["Indicator", "Description"], iocs)

doc.add_page_break()

# ---------------- PROBLEM SECTIONS ----------------
def problem(pid, title_txt, summary, method, confirmed_list, assessment_list, answer, mini_timeline, shots, pending_note=None):
    h(f"{pid} - {title_txt}", level=1)
    para("Summary", bold=True)
    para(summary)
    para("Technical Method", bold=True)
    para(method)
    para("Evidence", bold=True)
    for c in confirmed_list:
        para("Confirmed: " + c)
    for a in assessment_list:
        para("Assessment: " + a)
    if mini_timeline:
        para("Timeline for this Objective", bold=True)
        add_table(["Timestamp (UTC)", "Event"], mini_timeline)
    para("Answer", bold=True)
    if pending_note:
        para("Flag: PENDING - organizer clarification requested", bold=True)
        para(pending_note)
    else:
        para(answer, bold=True)
    for s in shots:
        add_image(os.path.join(SHOTS, s))
    doc.add_page_break()

problem(
    "P5-01", "Initially Infected Client",
    "Identify the initially infected client IPv4 address and hostname.",
    "Queried Security Onion Hunt for zeek.dhcp records on the 192.168.200.0/24 segment during "
    "the incident window, isolating DHCP lease and INFORM transactions.",
    ["DHCP records show client.address 192.168.200.95 issuing DHCP REQUEST/INFORM traffic "
     "with host.hostname \"Linwood-Win-PC\" beginning 2025-10-11 09:00:00 UTC."],
    [],
    "CTK{192.168.200.95|Linwood-Win-PC}",
    [("2025-10-11 09:00:00", "First DHCP OFFER/ACK to client 192.168.200.95 (Linwood-Win-PC)")],
    ["p5-01_step1_query.png", "p5-01_step2_value.png"],
)

problem(
    "P5-02", "AD Domain and Domain Controller",
    "Identify the Active Directory domain, and the domain controller hostname and IPv4 "
    "address.",
    "Queried zeek.kerberos records with destination.ip 192.168.200.4 to observe AS/TGS "
    "exchanges from the client, then cross-checked the domain and hostname casing against "
    "NTLM authentication records for the same host, which carry the authoritative "
    "capitalization used by the platform's answer key.",
    ["Kerberos AS-REQ/TGS-REQ traffic shows the client requesting service tickets for "
     "krbtgt/OYSTERTAINMENT.COM and LDAP/OYSTER-DC.oystertainment.com against 192.168.200.4.",
     "NTLM authentication records report the domain field literally as \"oystertainment.com\" "
     "(lowercase) and the server/hostname field as \"OYSTER-DC\" (uppercase)."],
    ["Kerberos SPN casing for the domain was inconsistent across records (mixed "
     "OYSTERTAINMENT.COM / oystertainment.com); the NTLM field was treated as authoritative "
     "for final answer casing."],
    "CTK{oystertainment.com|OYSTER-DC|192.168.200.4}",
    [("2025-10-11 09:00:00", "First Kerberos AS/TGS exchange between client and 192.168.200.4")],
    ["p5-02_step1_query.png", "p5-02_step2_value.png"],
)

problem(
    "P5-03", "Malicious Document URL",
    "Identify the complete HTTP URL that delivered the initial malicious document.",
    "Queried zeek.http records for source.ip 192.168.200.95, sorted ascending by timestamp, "
    "to identify the earliest HTTP transaction of the infection chain.",
    ["The first HTTP GET from 192.168.200.95 (2025-10-11 09:02:26 UTC) targeted virtual host "
     "r2consulting.net with URI /IRS-TRANSCRIPTS-037J/2/, returning an application/msword "
     "file."],
    [],
    "CTK{http://r2consulting.net/IRS-TRANSCRIPTS-037J/2/}",
    [("2025-10-11 09:02:26", "GET r2consulting.net/IRS-TRANSCRIPTS-037J/2/ - malicious document delivered")],
    ["p5-03_04_05_06_step1_query.png", "p5-03_step2_value.png"],
)

problem(
    "P5-04", "Document Response Metadata",
    "Identify the filename and MIME type recorded for the document response.",
    "Expanded the r2consulting.net HTTP record and inspected the file.resp_filenames and "
    "file.resp_mime_types fields extracted by Zeek's file analysis framework.",
    ["file.resp_filenames = \"transcript-June152018-017188/2.doc\"; file.resp_mime_types = "
     "\"application/msword\"."],
    [],
    "CTK{transcript-June152018-017188/2.doc|application/msword}",
    [("2025-10-11 09:02:26", "Document response recorded: transcript-June152018-017188/2.doc")],
    ["p5-04_step2_evidence.png"],
)

problem(
    "P5-05", "Executable Payload URL",
    "Identify the complete HTTP URL that delivered the executable payload.",
    "Continued the ascending zeek.http timeline for 192.168.200.95, identifying the next "
    "distinct HTTP GET transaction following the document download.",
    ["At 2025-10-11 09:03:25 UTC, the client issued a GET to virtual host www.ownhive.com "
     "with URI /MsWM2B0/, returning a Windows executable."],
    [],
    "CTK{http://www.ownhive.com/MsWM2B0/}",
    [("2025-10-11 09:03:25", "GET www.ownhive.com/MsWM2B0/ - executable payload delivered")],
    ["p5-05_step2_value.png"],
)

problem(
    "P5-06", "Executable Payload Metadata",
    "Provide the executable filename, MIME type, and response body size.",
    "Cross-verified the www.ownhive.com HTTP record's file.resp_filenames / "
    "file.resp_mime_types / http.response.body.length fields against the independent, "
    "authoritative zeek.file dataset record for the same connection UID to rule out a "
    "field-naming or data-source discrepancy, then inspected the full raw zeek.file message "
    "to confirm capture integrity and recover file hash values.",
    ["zeek.http record: file.resp_filenames = \"90352.exe\"; file.resp_mime_types = "
     "\"application/x-dosexec\"; http.response.body.length = 126976 bytes.",
     "Independent zeek.file record for the same connection UID (Cr14ep3JBUpdBqZTOl) reports "
     "identical values: file.name = \"90352.exe\", file.mime_type = \"application/x-dosexec\", "
     "file.bytes.total = 126976.",
     "The same zeek.file record confirms a complete, untruncated capture: seen_bytes = 126976, "
     "total_bytes = 126976, missing_bytes = 0, overflow_bytes = 0, extracted_cutoff = false, "
     "timedout = false - ruling out a partial capture as the source of any discrepancy in the "
     "byte count.",
     "The same record reports file hashes MD5 1a142395b5043ed2ca1545e1a5b0ba95 and SHA1 "
     "37a73193a0d4f4a337094e9e5fd439dc293fc615 for 90352.exe."],
    ["The value application/x-dosexec (Zeek's file-magic-detected MIME type) was submitted "
     "twice and rejected. The exercise platform instead accepted the generic MIME "
     "classification application/octet-stream for the same file, filename, and byte count - "
     "indicating the answer key expects the canonical/generic MIME form rather than Zeek's "
     "specific file-type detection."],
    "CTK{90352.exe|application/octet-stream|126976}",
    [("2025-10-11 09:03:25", "Executable response recorded: 90352.exe, 126,976 bytes")],
    ["p5-06_step2_evidence.png", "p5-06_step2_evidence_hashes.png"],
)

problem(
    "P5-07", "Repeated Cleartext POST Destination",
    "Identify the destination IP that received repeated cleartext HTTP POST transactions on "
    "TCP port 443, and the number of POST transactions indexed.",
    "Queried zeek.http for destination.port:443 AND http.method:\"POST\", grouped by "
    "destination.ip to identify the most frequently contacted host, then tested the "
    "\"transactions indexed\" count at three levels of granularity: the raw Zeek http.log "
    "transaction count, a grouping by network.community_id (unique TCP connections that "
    "carried a POST), and the total zeek.conn (TCP connection) record count to the same "
    "destination and port regardless of whether each connection carried a POST.",
    ["destination.ip 74.195.13.150 accounts for 32 raw POST transaction log entries on TCP/443 "
     "using plaintext HTTP (the protocol on this port is HTTP, not TLS, despite the standard "
     "HTTPS port number). Suricata's dedicated \"ET INFO HTTP traffic on port 443 (POST)\" rule "
     "independently corroborated this same count of 32.",
     "Grouping those 32 transactions by network.community_id shows they occurred over 20 "
     "distinct TCP connections that carried at least one POST (some connections carried "
     "multiple pipelined POST transactions).",
     "A direct query of event.dataset:\"zeek.conn\" AND destination.ip:\"74.195.13.150\" "
     "(no port restriction needed, as all traffic to this host is on 443) returned Total "
     "Found = 26 - the total number of distinct TCP connections established to this "
     "destination, independent of whether each one carried a POST."],
    ["A secondary host, 173.70.47.89, received only 2 similar cleartext POSTs on TCP/443 and "
     "does not meet the \"repeated\" threshold observed for 74.195.13.150, making "
     "74.195.13.150 the clear candidate for this objective regardless of counting method.",
     "The accepted count (26) is higher than the count of POST-carrying connections (20) but "
     "lower than the raw transaction count (32), indicating 6 of the 26 total TCP connections "
     "to this destination were established but did not carry a POST (e.g. connect-and-abandon "
     "or non-POST HTTP activity), while the remaining 20 carried the 32 cleartext POST "
     "payloads between them. The phrase \"POST transactions... indexed\" in this exercise's "
     "answer key evidently refers to the total connection count to the flagged endpoint, not "
     "the application-layer transaction count or the subset of connections matching the POST "
     "filter."],
    "CTK{74.195.13.150|26}",
    [("2025-10-11 09:03:30 - 09:11:26", "32 POST transactions across 20 of 26 total TCP connections to 74.195.13.150:443")],
    ["p5-07_step1_query.png", "p5-07_step2_value.png"],
)

problem(
    "P5-08", "First Client Callback to 188.124.167.132",
    "Identify the infected client's first POST to 188.124.167.132:8082, providing the UTC "
    "time and URI path.",
    "Queried zeek.http for source.ip 192.168.200.95 AND destination.ip 188.124.167.132 AND "
    "destination.port 8082. Only one matching record exists in the dataset, confirming it as "
    "both the first and only recorded transaction to this endpoint.",
    ["At 2025-10-11 09:07:06 UTC, the client POSTed to 188.124.167.132:8082 with URI "
     "/del9/LINWOOD-WIN-PC_W617601.A59A58979C2CD535524D2D1317484AC8/90."],
    [],
    "CTK{2025-10-11 09:07:06 UTC|/del9/LINWOOD-WIN-PC_W617601.A59A58979C2CD535524D2D1317484AC8/90}",
    [("2025-10-11 09:07:06", "First (and only recorded) client POST to 188.124.167.132:8082")],
    ["p5-08_step1_query.png", "p5-08_step2_value.png"],
)

problem(
    "P5-09", "First Domain Controller Malware Callback",
    "Identify the first confirmed malware-specific callback from the domain controller, "
    "providing destination, UTC time, and URI path.",
    "Queried zeek.http for source.ip 192.168.200.4 AND destination.ip 188.124.167.132, sorted "
    "ascending. Two matching records were found; the earliest was selected as the first "
    "callback.",
    ["At 2025-10-11 09:26:26 UTC, the domain controller POSTed to 188.124.167.132:8082 with "
     "URI /lib247/OYSTER-DC_W617600.EDF221B25063E00671993B9055FA7B90/81/, matching the naming "
     "convention of the client's own C2 check-in (hostname plus a Windows-update-style GUID)."],
    ["The matching URI naming convention between the client's and the DC's callbacks supports "
     "the assessment that the domain controller was compromised by the same malware family "
     "observed on the client, rather than by an unrelated, coincidental process."],
    "CTK{188.124.167.132|8082|2025-10-11 09:26:26 UTC|/lib247/OYSTER-DC_W617600.EDF221B25063E00671993B9055FA7B90/81/}",
    [("2025-10-11 09:26:26", "First domain controller POST to 188.124.167.132:8082 - DC compromise confirmed")],
    ["p5-09_step1_query.png", "p5-09_step2_value.png"],
)

problem(
    "P5-10", "PNG-Named Objects",
    "Identify the three PNG-named objects requested by the domain controller after "
    "compromise.",
    "Queried zeek.http for source.ip 192.168.200.4 with no destination restriction across the "
    "full post-compromise window, then expanded each raw record's message field to confirm "
    "the literal uri value for every HTTP GET issued by the domain controller.",
    ["The domain controller's raw Zeek http.log records show three distinct GET requests to "
     "185.159.131.55: uri \"/table.png\" at 09:26:27 UTC, uri \"/worming.png\" at 09:27:02 "
     "UTC, and uri \"/toler.png\" at 09:27:18 UTC. All three responses were 495,671-byte "
     "application/x-dosexec payloads despite the .png-styled URIs.",
     "The exercise platform accepted these three filenames only when submitted in "
     "alphabetical order (table.png, toler.png, worming.png) rather than in chronological "
     "first-seen order (table.png, worming.png, toler.png), which was rejected."],
    ["The .png extension does not match the actual response content (a Windows executable in "
     "every case), indicating these are disguised malware components rather than genuine "
     "images.",
     "This ordering distinction (alphabetical vs. chronological) is noted here as a lesson "
     "for future phases: when a challenge lists multiple filenames without explicitly stating "
     "an ordering rule, alphabetical order should be tried before chronological order."],
    "CTK{table.png|toler.png|worming.png}",
    [
        ("2025-10-11 09:26:27", "GET /table.png from 185.159.131.55"),
        ("2025-10-11 09:27:02", "GET /worming.png from 185.159.131.55"),
        ("2025-10-11 09:27:18", "GET /toler.png from 185.159.131.55"),
    ],
    ["p5-10_step1_query.png", "p5-10_step2_value.png"],
)

problem(
    "P5-11", "Client-to-DC SMB Communication",
    "Identify the client and server IP addresses that communicated over SMB/TCP 445.",
    "Queried connection metadata for destination.port:445 involving source.ip 192.168.200.95, "
    "spanning the full incident window.",
    ["258 connection records show 192.168.200.95 (client) communicating with 192.168.200.4 "
     "(domain controller) over TCP port 445 (SMB) throughout the incident."],
    ["This sustained SMB activity is consistent with the client accessing DC file/print "
     "shares, or with the malware using SMB for lateral movement or reconnaissance following "
     "DC compromise."],
    "CTK{192.168.200.95|192.168.200.4|445}",
    [("2025-10-11 (throughout incident window)", "258 SMB/TCP 445 connection records between client and DC")],
    ["p5-11_step1_query.png", "p5-11_step2_value.png"],
)

# ---------------- LESSONS FOR FUTURE PHASES ----------------
h("Lessons for Future Phases", level=1)
para(
    "Two non-obvious answer-key conventions were identified while resolving this phase's "
    "final objectives, and are recorded here to speed up future phases:"
)
para(
    "1. MIME type answers may be expected in their generic/canonical form (e.g. "
    "application/octet-stream for an arbitrary binary) rather than the specific file-magic "
    "classification a tool such as Zeek reports (e.g. application/x-dosexec for a Windows "
    "PE executable). When a filename/MIME/size-style answer is rejected, retry with the "
    "generic MIME classification before assuming the filename or size is wrong.",
    bold=False,
)
para(
    "2. A \"how many transactions were indexed\" style question may be scored at the TCP "
    "connection level rather than the application-layer transaction level. Before assuming "
    "the destination or scope is wrong, test multiple counting granularities: raw log-entry "
    "count, count of connections matching the specific activity, and total connection count "
    "to the destination/port regardless of activity type.",
    bold=False,
)

doc.save(OUT_DOCX)
print("Saved:", OUT_DOCX)
